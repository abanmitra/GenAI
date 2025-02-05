from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # ✅ Correct

def write_to_msword(file_path, generated_text):
    try:
        # Create a new Word document
        doc = Document()

        # Add a title with formatting
        title = doc.add_heading("Generated Response", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.style.font.size = Pt(16)  # Title font size

        # Split the generated text into paragraphs (split by newlines)
        paragraphs = generated_text.split('\n')

        # Add each paragraph to the document
        for paragraph in paragraphs:
            if paragraph.strip():  # Skip empty lines
                p = doc.add_paragraph()
                p.add_run(paragraph.strip()).font.size = Pt(12)  # Body text size

        # Optional: Add bullet points for lists (if your text has list items)
        # Example: Split by bullets (e.g., "•", "-", "*") and add as bullets
        if any(char in generated_text for char in ['•', '-', '*']):
            doc.add_paragraph("Key Points:", style='Heading 2')
            for line in generated_text.split('\n'):
                if line.strip().startswith(('•', '-', '*')):
                    doc.add_paragraph(line.strip().lstrip('•*-'), style='List Bullet')

        doc.save(file_path)

    except Exception as e:
        print(f"An error occurred: {e}")
        